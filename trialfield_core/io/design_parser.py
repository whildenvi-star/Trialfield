"""Parse PDF, docx, or prose input into a TrialDesign via the Anthropic API."""

from __future__ import annotations

import base64
import zipfile
from pathlib import Path
from typing import Optional

import anthropic

from trialfield_core.models.trial_design import TrialDesign, Treatment, TrialType

_SYSTEM_PROMPT = """You are an agricultural trial design parser. Extract trial design information \
from the provided text or image and call the create_trial_design tool with the structured result.

Rules:
- trial_type must be one of: fertility, seeding, spray, tillage, ground_speed, other
- For numeric treatments (nitrogen rates, seeding rates, etc.) set value to the numeric amount \
and unit to the appropriate unit string (e.g. "lb N/ac", "seeds/ac")
- For categorical treatments (product names, tank mixes, etc.) set value to null and unit to ""
- reps defaults to 4 if not specified
- plot_length_ft: use the value from the document if explicitly stated; otherwise set to null \
and the system will apply the trial-type default
- Extract treatment labels exactly as they appear; keep them concise
- If rates are given as "28k" or "28K", interpret as 28000
- Do not invent treatments not present in the input"""

_TOOL: dict = {
    "name": "create_trial_design",
    "description": "Create a structured trial design from the provided information.",
    "input_schema": {
        "type": "object",
        "properties": {
            "name": {
                "type": "string",
                "description": "Short descriptive name for the trial (e.g. 'N-rate trial')",
            },
            "trial_type": {
                "type": "string",
                "enum": ["fertility", "seeding", "spray", "tillage", "ground_speed", "other"],
            },
            "treatments": {
                "type": "array",
                "minItems": 2,
                "items": {
                    "type": "object",
                    "properties": {
                        "label": {"type": "string"},
                        "value": {
                            "type": ["number", "null"],
                            "description": "Numeric rate, or null for categorical treatments",
                        },
                        "unit": {"type": "string", "description": "e.g. 'lb N/ac', 'seeds/ac', ''"},
                    },
                    "required": ["label", "value", "unit"],
                },
            },
            "reps": {"type": "integer", "minimum": 2, "maximum": 8},
            "plot_length_ft": {
                "type": ["number", "null"],
                "description": "Plot length in feet; null to use trial-type default",
            },
        },
        "required": ["trial_type", "treatments", "reps"],
    },
}

_MODEL = "claude-sonnet-4-6"
_CLIENT: Optional[anthropic.Anthropic] = None


def _client() -> anthropic.Anthropic:
    global _CLIENT
    if _CLIENT is None:
        _CLIENT = anthropic.Anthropic()
    return _CLIENT


def _call_claude(content: list) -> TrialDesign:
    """Send content to Claude with forced tool use; return the validated TrialDesign."""
    response = _client().messages.create(
        model=_MODEL,
        max_tokens=1024,
        system=[
            {
                "type": "text",
                "text": _SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        tools=[_TOOL],
        tool_choice={"type": "tool", "name": "create_trial_design"},
        messages=[{"role": "user", "content": content}],
    )

    for block in response.content:
        if block.type == "tool_use" and block.name == "create_trial_design":
            raw = block.input
            treatments = [Treatment(**t) for t in raw["treatments"]]
            return TrialDesign(
                name=raw.get("name", "Untitled Trial"),
                trial_type=TrialType(raw["trial_type"]),
                treatments=treatments,
                reps=raw["reps"],
                plot_length_ft=raw.get("plot_length_ft"),
            )

    raise RuntimeError("Claude did not return a create_trial_design tool call")


def _extract_docx_content(path: Path) -> list:
    """Return a list of Claude content blocks (text + images) from a docx file."""
    from docx import Document

    content: list = []

    doc = Document(path)
    texts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))
    if texts:
        content.append({"type": "text", "text": "\n".join(texts)})

    _EXT_TO_MEDIA = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    with zipfile.ZipFile(path, "r") as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            ext = Path(name).suffix.lower()
            media_type = _EXT_TO_MEDIA.get(ext)
            if media_type is None:
                continue
            data = base64.standard_b64encode(z.read(name)).decode()
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": data},
            })

    if not content:
        raise ValueError(f"No extractable content found in {path}")

    return content


def _extract_pdf_content(path: Path) -> list:
    """Return a list of Claude content blocks (text) from a PDF file."""
    from pypdf import PdfReader

    reader = PdfReader(path)
    texts = [page.extract_text().strip() for page in reader.pages if page.extract_text().strip()]

    if not texts:
        raise ValueError(f"No extractable text found in {path}")

    return [{"type": "text", "text": "\n\n".join(texts)}]


def parse_prose(prose: str) -> TrialDesign:
    """Parse a free-text description of a trial into a TrialDesign."""
    return _call_claude([{"type": "text", "text": prose}])


def parse_docx(path: str | Path) -> TrialDesign:
    """Parse a Word document (.docx) into a TrialDesign."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    return _call_claude(_extract_docx_content(p))


def parse_pdf(path: str | Path) -> TrialDesign:
    """Parse a PDF document into a TrialDesign."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    return _call_claude(_extract_pdf_content(p))


def parse_design(
    prose: Optional[str] = None,
    docx_path: Optional[str | Path] = None,
    pdf_path: Optional[str | Path] = None,
) -> TrialDesign:
    """Parse a trial design from any combination of prose, docx, and/or pdf inputs."""
    if prose is None and docx_path is None and pdf_path is None:
        raise ValueError("At least one of prose, docx_path, or pdf_path must be provided")

    content: list = []

    if docx_path is not None:
        p = Path(docx_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {p}")
        content.extend(_extract_docx_content(p))

    if pdf_path is not None:
        p = Path(pdf_path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {p}")
        content.extend(_extract_pdf_content(p))

    if prose is not None:
        content.append({"type": "text", "text": prose})

    return _call_claude(content)
